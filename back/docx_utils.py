from pathlib import Path

from docx import Document


def get_docx_stats(file_path: Path) -> dict:
    doc = Document(file_path)

    return {
        "num_paragraphs": len(doc.paragraphs),
        "num_paragraphs_not_empty": len([p for p in doc.paragraphs if p.text]),
        "num_tables": len(doc.tables),
        "has_header": any(
            not section.header.is_linked_to_previous for section in doc.sections
        ),
        "has_footer": any(
            not section.footer.is_linked_to_previous for section in doc.sections
        ),
        "num_images": sum(
            1 for rel in doc.part.rels.values() if "image" in rel.target_ref
        ),
        "num_sections": len(doc.sections),
    }

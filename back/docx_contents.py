from pathlib import Path

import pandas as pd
from docx import Document
from utils import get_cell_text


def process_table(paragrahp, table, title, filename):
    expected_column_names = (
        "TIPO",
        "FILENAME",
        "NOMBRE",
        "MUNICIPIO",
        "ÁMBITO",
        "WEB",
        "DIRECCIÓN",
        "CATEGORIAS / DESCRIPCIÓN",
        "DIRIGIDO A",
        "TELEFONO",
    )
    paragrahp_text = paragrahp.text.strip()

    if paragrahp_text != title:
        msg = f"The title of the table is not the expected one:\n{paragrahp_text}\n{title}"

        ignore_errors = filename in {
            "METROPOLITANA/GÁDOR.docx",
            "METROPOLITANA/RIOJA.docx",
            "FILABRES/TABERNAS.docx",
            "ALPUJARRA/ALBOLODUY.docx",
            "ALPUJARRA/FONDÓN.docx",
        }

        if ignore_errors:
            print(msg)
            print("Ignoring error")
        else:
            raise ValueError(msg)

    data = [[get_cell_text(cell) for cell in row.cells] for row in table.rows]
    data = [row for row in data if any(cell is not None for cell in row)]

    column_names = tuple(data[0])

    if column_names != expected_column_names[2:]:
        msg = f"Table does not have the correct column names. Or missing header:\n{column_names}\n{expected_column_names[2:]}"
        if (
            filename == "FILABRES/LAS TRES VILLAS.docx"
            and data[0][0] == "PISCINA DE OCAÑA"
        ):
            print(msg)
            print("Ignoring error")
            data.insert(0, expected_column_names[2:])
            column_names = expected_column_names[2:]
        elif filename in {
            "ALMANZORA/LÚCAR.docx",
            "ALMANZORA/ALBANCHEZ.docx",
            "ALPUJARRA/BEIRES.docx",
            "ALPUJARRA/ÍLLAR.docx",
            "ALPUJARRA/ALCOLEA.docx",
            "ALPUJARRA/ALHABIA.docx",
            "ALMANZORA/SIERRO.docx",
            "PONIENTE/ENIX.docx",
        }:
            print(msg)
            print("Ignoring error")
            data[0] = list(expected_column_names[2:])
            column_names = tuple(data[0])
        elif (filename == "ALPUJARRA/BENTARIQUE.docx") and (
            set(expected_column_names[2:]) - set(column_names) == {"DIRIGIDO A"}
        ):
            print(msg)
            print("Ignoring error")
            insert_index = expected_column_names[2:].index("DIRIGIDO A")
            data[0].insert(insert_index, "DIRIGIDO A")
            for row in data[1:]:
                row.insert(insert_index, None)
            column_names = tuple(data[0])
        elif (filename == "ALPUJARRA/ALMÓCITA.docx") and (
            set(expected_column_names[2:]) - set(column_names)
            == {"WEB", "CATEGORIAS / DESCRIPCIÓN"}
        ):
            print(msg)
            print("Ignoring error")
            insert_index = expected_column_names[2:].index("WEB")
            data[0].insert(insert_index, "WEB")
            for row in data[1:]:
                row.insert(insert_index, None)
            insert_index = expected_column_names[2:].index("CATEGORIAS / DESCRIPCIÓN")
            data[0][insert_index] = "CATEGORIAS / DESCRIPCIÓN"
            column_names = tuple(data[0])

        else:
            raise ValueError(msg)
    df = pd.DataFrame(data[1:], columns=column_names)

    df.insert(0, expected_column_names[0], (title,) * len(df))
    df.insert(1, expected_column_names[1], (filename,) * len(df))

    df.loc[df["FILENAME"] == "ALPUJARRA/ALBOLODUY.docx", "MUNICIPIO"] = "Alboloduy"
    df.loc[df["FILENAME"] == "ALPUJARRA/FONDÓN.docx", "MUNICIPIO"] = "Fondón"
    df.loc[df["FILENAME"] == "ALPUJARRA/ALCOLEA.docx", "MUNICIPIO"] = "Alcolea"
    df.loc[df["FILENAME"] == "ALPUJARRA/BENTARIQUE.docx", "MUNICIPIO"] = "Bentarique"
    df.loc[df["FILENAME"] == "ALMANZORA/PURCHENA.docx", "MUNICIPIO"] = "Purchena"
    df.loc[df["FILENAME"] == "ALMANZORA/LAROYA.docx", "MUNICIPIO"] = "Laroya"
    df.loc[df["FILENAME"] == "ALMANZORA/LÍJAR.docx", "MUNICIPIO"] = "Líjar"
    df.loc[df["FILENAME"] == "PONIENTE/ENIX.docx", "MUNICIPIO"] = "Enix"
    df.loc[df["FILENAME"] == "FILABRES/CASTRO FILABRES.docx", "MUNICIPIO"] = (
        "Castro de Filabres"
    )
    df.loc[df["FILENAME"] == "FILABRES/BENIZALON.docx", "MUNICIPIO"] = "Benizalón"

    df["MUNICIPIO"] = (
        df["MUNICIPIO"]
        .replace("Santa Fé de Mondújar", "Santa Fe de Mondújar")
        .replace("PECHINA", "Pechina")
        .replace("las Tres Villas", "Las Tres Villas")
        .replace("fiñana", "Fiñana")
        .replace("Nacimiento “", "Nacimiento")
        .replace("ALSODUX", "Alsodux")
        .replace("Sufli", "Suflí")
        .replace("Aboleas", "Arboleas")
    )

    if df["MUNICIPIO"].nunique(dropna=False) != 1:
        raise ValueError(
            f"The 'MUNICIPIO' column must have the same value for all rows in the dataframe. Found: {df['MUNICIPIO'].unique()}"
        )

    df["DIRIGIDO A"] = (
        df["DIRIGIDO A"]
        .replace(" y ", " ", regex=True)
        .str.replace(r"[,\n/]+", " ", regex=True)
        .str.replace(r"[ ]+", " ", regex=True)
        .str.replace("ciudadano", "Ciudadanía", case=False)
        .str.replace("emrpesa", "Empresa", case=False)
        .str.replace("empresas", "Empresa", case=False)
        .str.replace("ciuddano", "Ciudadanía", case=False)
        .str.replace("ciudadanías", "Ciudadanía", case=False)
        .str.replace("ciudadanía.", "Ciudadanía", case=False)
        .str.replace("ciudadanía empresa", "Ambas", case=False)
        .str.replace("empresa ciudadanía", "Ambas", case=False)
        .str.replace("ambos", "Ambas", case=False)
    )

    return df


def process_docx(file_path: Path) -> pd.DataFrame:
    first_table_title = "Entidades o servicios interés municipal"
    second_table_title = "Lugares de interés municipal"
    filename = "/".join(file_path.parts[-2:])
    doc = Document(file_path)

    if filename in {
        "ALPUJARRA/CANJAYAR.docx",
        "PONIENTE/ENIX.docx",
        "FILABRES/CASTRO FILABRES.docx",
    }:
        doc.paragraphs[0].text = first_table_title
        return process_table(
            doc.paragraphs[0], doc.tables[0], first_table_title, filename
        )
    not_empty_paragraph_indexes = [i for i, p in enumerate(doc.paragraphs) if p.text]
    first_not_empty_paragraph = doc.paragraphs[not_empty_paragraph_indexes[0]]
    second_not_empty_paragraph = doc.paragraphs[not_empty_paragraph_indexes[1]]
    df0 = process_table(
        first_not_empty_paragraph, doc.tables[0], first_table_title, filename
    )
    if filename == "FILABRES/TAHAL.docx":
        return df0
    df1 = process_table(
        second_not_empty_paragraph, doc.tables[1], second_table_title, filename
    )

    return pd.concat([df0, df1])
